from docx import Document
from docx.shared import RGBColor

doc = Document()
doc.add_heading("파이썬으로 워드파일 생성하기", level=1)
p = doc.add_paragraph("파이썬으로 작성하는 워드의 문장입니다.")
p.add_run("add_paragraph()가 반환한 객체를 통해 ")
p.add_run("추가 문장을 삽입할 수 있습니다.")
p.add_run("글자의 색을 변경하거나").font.color.rgb = RGBColor(255, 0, 255)
p.add_run("굵게 표기하거나, ").bold=True
p.add_run("이탤릭 효과도 줍니다.").italic = True
doc.add_paragraph("python-docx라이브러리를 사용합니다.")
doc.add_paragraph("라이브러리에는 수많은 기능이 구현되어 있습니다.")
doc.save("docx_sample2.docx")