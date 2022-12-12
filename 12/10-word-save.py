from docx import Document
from docx.shared import Inches

headers = ["이름", "국어", "영어", "수학", "음악", "미술", "체육", "컴퓨터"]
datas = [
    ["홍길동", 100, 99, 88, 77, 66, 77, 44],
    ["김길동", 95, 93, 85, 67, 86, 63, 84],
    ["최길동", 77, 46, 78, 89, 93, 78, 98],
    ["박길동", 96, 94, 96, 95, 91, 91, 85],
    ["차길동", 87, 91, 89, 99, 95, 94, 92],
    ["구길동", 78, 83, 88, 65, 58, 72, 68],
]

doc = Document()
doc.add_heading("파이썬으로 워드파일 생성하기", level=1)
doc.add_paragraph("파이썬으로 작성하는 워드의 문장입니다.")
doc.add_paragraph("python-docx라이브러리를 사용합니다.")
doc.add_paragraph("라이브러리에는 수많은 기능이 구현되어 있습니다.")

doc.add_picture("./12/sample.jpg", width=Inches(5.0))

table = doc.add_table(rows=1, cols=8)
table.style = "Table Grid"
header = table.rows[0].cells
for i, h in enumerate(headers):
    header[i].text = h
for data in datas:
    row = table.add_row().cells
    for i, d in enumerate(data):
        row[i].text = str(d)
doc.save("docx_sample.docx")