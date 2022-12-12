from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt

doc = Document()
my_style = doc.styles
font_style = my_style.add_style("st", WD_STYLE_TYPE.CHARACTER)
font = font_style.font
font.size = Pt(15)

p2 = doc.add_paragraph("파이썬으로 워드문서를 작성!! ")
p2.add_run("스타일을 적용하는 방법도 있습니다. ", style="st")
p2.add_run("스타일과 속성을 혼합 사용", style="st").bold = True
doc.save("docx_sample3.docx")